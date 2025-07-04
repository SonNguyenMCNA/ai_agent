
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from openai import OpenAI

client = OpenAI(api_key=st.secrets["key"])

st.set_page_config(page_title="AI Agent Báo Cáo Đào Tạo", layout="centered")
st.title("📊 AI Agent – Tổng Hợp Báo Cáo Đào Tạo Tự Động")
st.markdown("Tải lên 3 file Excel và 1 file Word template để tạo báo cáo tự động.")

hoc_vien_file = st.file_uploader("📘 1. Danh sách học viên (HocVien.xlsx)", type=["xlsx"])
diem_danh_file = st.file_uploader("📝 2. Danh sách điểm danh (DiemDanh.xlsx)", type=["xlsx"])
ket_qua_file = st.file_uploader("📈 3. Kết quả cuối khóa (KetQua.xlsx)", type=["xlsx"])
template_file = st.file_uploader("📄 4. File mẫu báo cáo (Word Template)", type=["docx"])

def set_paragraph_format(paragraph, font_name="Arial", font_size=12, spacing=1.3):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = False
    paragraph.style.font.name = font_name
    paragraph.style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    spacing_elm = OxmlElement("w:spacing")
    spacing_elm.set(qn("w:line"), str(int(spacing * 240)))
    spacing_elm.set(qn("w:lineRule"), "auto")
    pPr.append(spacing_elm)

if st.button("🚀 Tạo báo cáo") and all([hoc_vien_file, diem_danh_file, ket_qua_file, template_file]):
    try:
        df_hoc_vien = pd.read_excel(hoc_vien_file)
        df_diem_danh = pd.read_excel(diem_danh_file)
        df_ket_qua = pd.read_excel(ket_qua_file)

        total_students = len(df_hoc_vien)
        df_diem_danh['Số buổi tham gia'] = df_diem_danh.iloc[:, 1:-1].apply(lambda x: x.eq('X').sum(), axis=1)
        attendance_rate = round(df_diem_danh['Số buổi tham gia'].mean() / (df_diem_danh.shape[1] - 2) * 100, 2)
        completed_students = df_ket_qua[df_ket_qua['Tổng điểm'] >= 7]
        completion_rate = round(len(completed_students) / total_students * 100, 2)
        top_students = df_ket_qua.sort_values(by='Tổng điểm', ascending=False).head(3)
        vang_phep = df_diem_danh['Ghi chú'].str.contains("có phép", case=False).sum()
        gioi_xuat_sac_rate = round(len(df_ket_qua[df_ket_qua['Xếp loại'].isin(['Giỏi', 'Xuất sắc'])]) / total_students * 100, 2)

        prompt = f"""
Bạn đóng vai trò là hệ thống đánh giá đào tạo nội bộ tại một doanh nghiệp lớn (ví dụ: Viettel). 
Hãy viết một đoạn nhận xét từ 4–6 câu, đánh giá tổng quan khóa học dựa trên các thông tin sau:

- Tổng số học viên: {total_students}
- Tỉ lệ hoàn thành khóa học: {completion_rate}%
- Tỉ lệ đạt loại Giỏi – Xuất sắc: {gioi_xuat_sac_rate}%
- Tỉ lệ tham gia điểm danh trung bình: {attendance_rate}%
- Số học viên vắng có phép: {vang_phep}
- 3 học viên điểm cao nhất: {top_students['Họ tên'].tolist()} với điểm {top_students['Tổng điểm'].tolist()}

Yêu cầu:
- Dùng giọng văn khách quan, chuyên nghiệp
- Nêu rõ xu hướng học tập
- Đánh giá năng lực chung
- Đề xuất ý tưởng/khuyến nghị nếu phù hợp

Kết quả trả về: một đoạn văn hoàn chỉnh.
"""
        try:
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}]
            )
            ai_comment = response.choices[0].message.content.strip()
        except Exception as e:
            ai_comment = "Không thể kết nối GPT để sinh nhận xét."
            st.error(f"Lỗi khi gọi GPT: {e}")

        doc = Document(template_file)
        placeholders = {
            "Khóa học: ....................................................": "Khóa học: Ứng dụng AI vào công việc tại Viettel",
            "Thời gian: ....................................................": "Thời gian: 15–17/05/2025",
            "Số học viên: ........ người": f"Số học viên: {total_students} người",
            "Tỷ lệ hoàn thành: ........%": f"Tỷ lệ hoàn thành: {completion_rate}%",
            "Tỷ lệ đạt loại Giỏi – Xuất sắc: ........%": f"Tỷ lệ đạt loại Giỏi – Xuất sắc: {gioi_xuat_sac_rate}%",
            "- ....................................................": [f"- {row['Họ tên']} – {row['Tổng điểm']} điểm – {row['Xếp loại']}" for _, row in top_students.iterrows()],
            "- Trung bình mỗi học viên tham gia ........% số buổi": f"- Trung bình mỗi học viên tham gia {attendance_rate}% số buổi",
            "- Số trường hợp vắng mặt có phép: ...": f"- Số trường hợp vắng mặt có phép: {vang_phep}",
            "- ..............................................................................": [f"- {line.strip()}" for line in ai_comment.split(". ") if line.strip()]
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if text in placeholders:
                replacement = placeholders[text]
                if isinstance(replacement, list):
                    parent = para._element.getparent()
                    idx = parent.index(para._element)
                    parent.remove(para._element)
                    for i, val in enumerate(replacement):
                        new_para = doc.add_paragraph(val)
                        set_paragraph_format(new_para)
                        parent.insert(idx + i, new_para._element)
                else:
                    para.text = replacement
                    set_paragraph_format(para)

        output_stream = BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)

        st.success("🎉 Báo cáo đã được tạo thành công!")
        st.download_button(label="📥 Tải file báo cáo Word",
                           data=output_stream,
                           file_name="BaoCaoTongHop.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        st.error(f"Đã xảy ra lỗi khi xử lý: {e}")
